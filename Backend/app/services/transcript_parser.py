import docx
import re
from typing import List, Dict, Any, Tuple
from datetime import datetime
import uuid
import logging
from sqlalchemy.ext.asyncio import AsyncSession
from app.models.database import ConversationLog, Customer
from app.core.qdrant_client import QdrantVectorDB
from app.services.embedding_service import EmbeddingService
from app.utils.dates import session_to_ymd

logger = logging.getLogger(__name__)

class TranscriptParserService:
    """Parses transcripts and populates Store 1 (Full Conversation Log)"""
    
    def __init__(self):
        self.qdrant = QdrantVectorDB()
        self.embedding_service = EmbeddingService()
        
    async def parse_and_store_transcript(
        self, 
        file_path: str, 
        db: AsyncSession, 
        customer_id: uuid.UUID,
        transcript_id: uuid.UUID,
        session_name: str,
        call_date: datetime,
        client_speaker_name: str
    ) -> List[Dict[str, Any]]:
        """
        Parse DOCX transcript, store in Postgres (ConversationLog) and Qdrant (conversations).
        Returns a list of parsed conversation blocks.
        """
        logger.info(f"Parsing transcript: {file_path} for session {session_name}")
        
        try:
            doc = docx.Document(file_path)
        except Exception as e:
            logger.error(f"Failed to read DOCX file {file_path}: {e}")
            raise
            
        blocks = self._extract_speaker_blocks(doc)
        logger.info(f"Extracted {len(blocks)} conversation blocks from transcript")
        
        # Determine roles
        for block in blocks:
            # Normalize whitespace: "Prasad   Kadrikar" -> "prasad kadrikar"
            speaker_norm = " ".join(block['speaker'].split()).lower()
            client_norm = " ".join(client_speaker_name.split()).lower() if client_speaker_name else ""
            
            if client_norm and speaker_norm == client_norm:
                block['role'] = 'client'
            else:
                block['role'] = 'team_member'
                
            block['session'] = session_name
            block['call_date'] = call_date.isoformat()
            block['customer_id'] = str(customer_id)
            
        # 1. Save to PostgreSQL (Store 1)
        db_logs = []
        # Use the transcript_id passed in from the caller (same ID as in transcripts table)
        
        for block in blocks:
            log_entry = ConversationLog(
                id=uuid.uuid4(),
                transcript_id=transcript_id,
                customer_id=customer_id,
                speaker=block['speaker'],
                role=block['role'],
                text=block['text'],
                call_timestamp=block['timestamp']
            )
            db.add(log_entry)
            db_logs.append(log_entry)
            
            # Keep ID in block for Qdrant payload linking
            block['log_id'] = str(log_entry.id)
            block['transcript_id'] = str(transcript_id)
            
        await db.commit()
        logger.info("Saved conversation logs to PostgreSQL")
        
        # 2. Generate Embeddings & Save to Qdrant (Store 1)
        # Implement conversational sliding window (N=5: i-2 to i+2) to preserve context
        texts_to_embed = []
        for i in range(len(blocks)):
            start_idx = max(0, i - 2)
            end_idx = min(len(blocks), i + 3) # exclusive
            
            context_window = []
            for j in range(start_idx, end_idx):
                b = blocks[j]
                context_window.append(f"{b['speaker']} [{b['timestamp']}]: {b['text']}")
                
            combined_context = "\n".join(context_window)
            texts_to_embed.append(combined_context)
            
        embeddings = self.embedding_service.generate_embeddings(texts_to_embed)
        
        # Normalized filterable date (YYYY-MM-DD), derived from the session name so it is
        # consistent across both vector stores and immune to the timezone-shifted timestamp.
        date_ymd = session_to_ymd(session_name, call_date)

        # Prepare Qdrant payloads
        payloads = []
        ids = []
        for block in blocks:
            payloads.append({
                "log_id": block["log_id"],
                "transcript_id": block["transcript_id"],
                "customer_id": block["customer_id"],
                "speaker": block["speaker"],
                "role": block["role"],
                "text": block["text"],
                "call_timestamp": block["timestamp"],
                "session": block["session"],
                "call_date": block["call_date"],
                "date_ymd": date_ymd
            })
            ids.append(block["log_id"])
            
        success = self.qdrant.add_vectors(
            vectors=embeddings,
            payloads=payloads,
            ids=ids,
            collection_name=self.qdrant.conversations_collection
        )
        
        if not success:
            logger.error("Failed to add conversation vectors to Qdrant")
            
        return blocks
        
    def _extract_speaker_blocks(self, doc) -> List[Dict[str, str]]:
        """
        Extracts clean speaker blocks from the raw DOCX paragraphs.
        
        Handles the Microsoft Teams transcript format where each paragraph
        contains a full conversation block: speaker header + text joined by \\n.
        Example paragraph text:
            '\\nPrasad Kadrikar   0:04\\nYeah, let's get started.'
        """
        blocks = []
        
        # Pattern to match speaker header line: 'Speaker Name   1:23' or '01:23:45'
        speaker_pattern = re.compile(r'^([A-Za-z0-9 ]+\S)\s{2,}(\d+:\d+(?::\d+)?)$')
        
        for para in doc.paragraphs:
            raw = para.text
            if not raw.strip():
                continue
            
            # Split the paragraph into individual lines
            lines = [line.strip() for line in raw.split('\n')]
            lines = [l for l in lines if l]  # Remove empty lines
            
            if not lines:
                continue
            
            # Skip preamble paragraphs
            if any("Meeting Recording" in l or l.endswith("started transcription") for l in lines):
                continue
            
            # Check if the first line is a speaker header
            first_line = lines[0]
            match = speaker_pattern.match(first_line)
            
            if match:
                # This paragraph is a self-contained conversation block
                speaker = match.group(1).strip()
                timestamp = match.group(2).strip()
                text_lines = lines[1:]  # Everything after the header is the spoken text
                
                if text_lines:
                    combined_text = " ".join(text_lines).strip()
                    # Filter out noise (very short utterances like "Okay", "Yeah")
                    if len(combined_text.split()) > 2 or len(combined_text) > 10:
                        blocks.append({
                            "speaker": speaker,
                            "timestamp": timestamp,
                            "text": combined_text
                        })
            else:
                # Fallback: treat as plain continuation text (old format support)
                logger.debug(f"Skipping non-speaker paragraph: {first_line[:60]}")
                
        logger.info(f"Extracted {len(blocks)} blocks from {len(doc.paragraphs)} paragraphs")
        return blocks
